from document_translator import DocumentTranslator
from docx import Document
import os

def main():
    # Initialize the translator
    translator = DocumentTranslator()
    
    # Directory containing the files
    input_dir = "pptx_files"
    output_dir = "translated_files"
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # List of files to translate
    files_to_translate = [
        "acetatosILN_TL-aula4a.pdf",
        "acetatosILN_TL-aula4b.pdf"
    ]
    
    # Translate each file
    for filename in files_to_translate:
        input_path = os.path.join(input_dir, filename)
        output_path = os.path.join(output_dir, f"translated_{os.path.splitext(filename)[0]}.docx")
        
        print(f"\nTranslating {filename}...")
        try:
            # Extract text from PDF
            doc_content = translator.parsers[".pdf"].extract_content(input_path)
            
            # Translate the text
            translated_text = translator._translate_content(
                doc_content["text"],
                source_lang="pt",
                target_lang="en"
            )
            
            # Create Word document
            doc = Document()
            
            # Add title
            doc.add_heading(f"Translation of {filename}", 0)
            
            # Add translated content
            for paragraph in translated_text.split('\n'):
                if paragraph.strip():  # Only add non-empty paragraphs
                    doc.add_paragraph(paragraph)
            
            # Save the document
            doc.save(output_path)
            print(f"Successfully translated {filename} to {output_path}")
            
        except Exception as e:
            print(f"Error translating {filename}: {str(e)}")

if __name__ == "__main__":
    main() 