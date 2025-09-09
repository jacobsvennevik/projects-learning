from typing import Dict, Any
from docx import Document
import json

class DOCXParser:
    def extract_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and formatting information from a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dictionary containing text and formatting information
        """
        doc = Document(file_path)
        text_content = []
        formatting_info = []
        
        for paragraph in doc.paragraphs:
            paragraph_content = {
                "text": paragraph.text,
                "formatting": {
                    "style": paragraph.style.name,
                    "alignment": str(paragraph.alignment),
                    "runs": []
                }
            }
            
            for run in paragraph.runs:
                run_content = {
                    "text": run.text,
                    "formatting": {
                        "font_name": run.font.name,
                        "font_size": run.font.size.pt if run.font.size else None,
                        "is_bold": run.font.bold,
                        "is_italic": run.font.italic,
                        "is_underline": run.font.underline
                    }
                }
                paragraph_content["formatting"]["runs"].append(run_content)
                text_content.append(run.text)
            
            formatting_info.append(paragraph_content)
        
        return {
            "text": "\n".join(text_content),
            "formatting": formatting_info
        }
    
    def create_document(self, file_path: str, translated_text: str, 
                       formatting_info: Dict[str, Any]) -> None:
        """
        Create a new Word document with translated content and preserved formatting.
        
        Args:
            file_path: Path where the new Word document should be saved
            translated_text: The translated text content
            formatting_info: Dictionary containing formatting information
        """
        doc = Document()
        
        # Split translated text into chunks (one per paragraph)
        translated_chunks = translated_text.split("\n")
        chunk_index = 0
        
        for paragraph_info in formatting_info:
            if chunk_index < len(translated_chunks):
                # Create new paragraph with same style
                paragraph = doc.add_paragraph(style=paragraph_info["formatting"]["style"])
                paragraph.alignment = paragraph_info["formatting"]["alignment"]
                
                # Add the translated text with preserved formatting
                run = paragraph.add_run(translated_chunks[chunk_index])
                
                # Apply formatting from the first run (since we're combining all runs)
                if paragraph_info["formatting"]["runs"]:
                    first_run = paragraph_info["formatting"]["runs"][0]
                    formatting = first_run["formatting"]
                    
                    if formatting["font_name"]:
                        run.font.name = formatting["font_name"]
                    if formatting["font_size"]:
                        run.font.size = formatting["font_size"]
                    run.font.bold = formatting["is_bold"]
                    run.font.italic = formatting["is_italic"]
                    run.font.underline = formatting["is_underline"]
                
                chunk_index += 1
        
        doc.save(file_path) 